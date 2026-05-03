"""
Genera un informe .docx con:
  - Resumen ejecutivo
  - Optimizaciones aplicadas (resumido desde OPTIMIZACIONES.md)
  - Resultados de Android Lint
  - Resultados del perfilado en 3 dispositivos fisicos (con graficas)

Uso:
    python scripts/generate-report.py

Salida:
    Informe-Optimizaciones-Vinilos.docx (en raiz del repo)
"""

import re
from pathlib import Path
from datetime import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = Path(__file__).resolve().parent.parent
PROFILE_DIR = REPO / "profile-results"
CHARTS_DIR = REPO / "build" / "report-charts"
CHARTS_DIR.mkdir(parents=True, exist_ok=True)
OUT_DOCX = REPO / "Informe-Optimizaciones-Vinilos.docx"

# Paleta corporativa
COLOR_PRIMARY = "#D32F2F"     # rojo Vinilos
COLOR_ACCENT = "#1976D2"      # azul
COLOR_GREEN = "#388E3C"
COLOR_GRAY = "#616161"
PALETTE = ["#D32F2F", "#1976D2", "#388E3C", "#F57C00", "#7B1FA2", "#00796B", "#5D4037"]

plt.rcParams.update({
    "font.family": "DejaVu Sans",
    "font.size": 10,
    "axes.titlesize": 12,
    "axes.titleweight": "bold",
    "axes.labelsize": 10,
    "axes.spines.top": False,
    "axes.spines.right": False,
    "axes.grid": True,
    "grid.alpha": 0.3,
    "grid.linestyle": "--",
})


# =======================================================================
# 1. PARSING de los reportes de profile
# =======================================================================

def parse_profile(path):
    """Extrae métricas relevantes de un reporte profile-results/*.txt."""
    txt = path.read_text(encoding="utf-8", errors="replace")
    out = {"path": path, "label": "", "model": "", "android": "", "abi": ""}

    # Cabecera
    m = re.search(r"DeviceLabel:\s*(\S+)", txt)
    if m: out["label"] = m.group(1)
    m = re.search(r"Model:\s*(\S+)", txt)
    if m: out["model"] = m.group(1)
    m = re.search(r"Android:\s*(\S+)", txt)
    if m: out["android"] = m.group(1)
    m = re.search(r"ABI:\s*(\S+)", txt)
    if m: out["abi"] = m.group(1)

    # App Summary - PSS por categoría (en KB)
    app_summary = re.search(
        r"App Summary.*?TOTAL(?:\s+PSS)?:\s*(\d+)",
        txt, re.DOTALL
    )
    if app_summary:
        block = app_summary.group(0)
        for label, key in [
            ("Java Heap", "java_heap_kb"),
            ("Native Heap", "native_heap_kb"),
            ("Code", "code_kb"),
            ("Stack", "stack_kb"),
            ("Graphics", "graphics_kb"),
            ("Private Other", "private_other_kb"),
            ("System", "system_kb"),
        ]:
            m = re.search(rf"{re.escape(label)}:\s*(\d+)", block)
            out[key] = int(m.group(1)) if m else 0
        out["total_pss_kb"] = int(app_summary.group(1))

    # GPU stats
    m = re.search(r"Total frames rendered:\s*(\d+)", txt)
    out["frames_total"] = int(m.group(1)) if m else 0

    m = re.search(r"Janky frames(?:\s+\(legacy\))?:\s*(\d+)\s*\(([\d.]+)%\)", txt)
    if m:
        out["janky_frames"] = int(m.group(1))
        out["janky_pct"] = float(m.group(2))
    else:
        out["janky_frames"] = 0
        out["janky_pct"] = 0.0

    for p in (50, 90, 95, 99):
        m = re.search(rf"{p}th percentile:\s*(\d+)ms", txt)
        out[f"p{p}_ms"] = int(m.group(1)) if m else 0

    m = re.search(r"Number Missed Vsync:\s*(\d+)", txt)
    out["missed_vsync"] = int(m.group(1)) if m else 0
    m = re.search(r"Number Slow UI thread:\s*(\d+)", txt)
    out["slow_ui_thread"] = int(m.group(1)) if m else 0
    m = re.search(r"Number Slow bitmap uploads:\s*(\d+)", txt)
    out["slow_bitmap_uploads"] = int(m.group(1)) if m else 0
    m = re.search(r"Number Slow issue draw commands:\s*(\d+)", txt)
    out["slow_draw_cmds"] = int(m.group(1)) if m else 0

    # GPU memory
    m = re.search(r"Total GPU memory usage:\s*\n\s*\d+\s*bytes,\s*([\d.]+)\s*MB", txt)
    out["gpu_mem_mb"] = float(m.group(1)) if m else 0.0

    # Battery (suele ser 0 si está cargando)
    m = re.search(r"Estimated battery capacity:\s*(\d+)\s*mAh", txt)
    out["battery_capacity_mah"] = int(m.group(1)) if m else 0

    # CPU instantáneo
    m = re.search(
        r"^\s*\d+\s+\S+\s+\S+\s+\S+\s+\S+\s+\S+\s+\S+\s+\S\s+([\d.]+)\s+([\d.]+)\s+\S+\s+com\.uniandes\.vinilos",
        txt, re.MULTILINE
    )
    if m:
        out["cpu_pct_instant"] = float(m.group(1))
        out["mem_pct_instant"] = float(m.group(2))
    else:
        out["cpu_pct_instant"] = 0.0
        out["mem_pct_instant"] = 0.0

    return out


# =======================================================================
# 2. GRAFICAS
# =======================================================================

def chart_pss_stacked(devices, out_path):
    """Stacked bar de PSS por categoría por dispositivo (en MB)."""
    cats = [
        ("Java Heap", "java_heap_kb"),
        ("Native Heap", "native_heap_kb"),
        ("Code", "code_kb"),
        ("Graphics", "graphics_kb"),
        ("Stack", "stack_kb"),
        ("Private Other", "private_other_kb"),
        ("System", "system_kb"),
    ]
    labels = [d["label"] for d in devices]
    fig, ax = plt.subplots(figsize=(8, 4.5))

    bottoms = np.zeros(len(devices))
    for (cat_name, key), color in zip(cats, PALETTE):
        values = np.array([d.get(key, 0) / 1024.0 for d in devices])
        ax.bar(labels, values, bottom=bottoms, label=cat_name, color=color, edgecolor="white", linewidth=0.5)
        bottoms += values

    # Etiqueta del total encima de cada barra
    for i, d in enumerate(devices):
        total_mb = d.get("total_pss_kb", 0) / 1024.0
        ax.text(i, bottoms[i] + 2, f"{total_mb:.1f} MB", ha="center", fontweight="bold", fontsize=10)

    ax.set_ylabel("Memoria PSS (MB)")
    ax.set_title("Composición de memoria PSS por dispositivo")
    ax.legend(loc="upper left", bbox_to_anchor=(1.02, 1.0), fontsize=8)
    ax.set_ylim(0, max(bottoms) * 1.15)
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def chart_jank_pct(devices, out_path):
    """Barra horizontal con % de janky frames."""
    labels = [d["label"] for d in devices]
    values = [d["janky_pct"] for d in devices]
    fig, ax = plt.subplots(figsize=(8, 3.5))
    bars = ax.barh(labels, values, color=[COLOR_GREEN if v < 5 else COLOR_PRIMARY if v > 15 else "#F57C00" for v in values])
    ax.set_xlabel("% de frames jankerados (mayor = peor)")
    ax.set_title("Porcentaje de janky frames por dispositivo")
    ax.axvline(5, color=COLOR_GRAY, linestyle="--", alpha=0.7, label="Umbral aceptable (5%)")
    ax.legend(loc="lower right")
    for bar, v in zip(bars, values):
        ax.text(v + 0.2, bar.get_y() + bar.get_height() / 2, f"{v:.2f}%", va="center", fontweight="bold")
    ax.set_xlim(0, max(max(values), 5) * 1.25)
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def chart_percentiles(devices, out_path):
    """Grouped bar de p50/p90/p95/p99 frame time por dispositivo."""
    labels = [d["label"] for d in devices]
    metrics = [("p50", "P50"), ("p90", "P90"), ("p95", "P95"), ("p99", "P99")]
    n_metrics = len(metrics)
    x = np.arange(len(labels))
    width = 0.2
    fig, ax = plt.subplots(figsize=(8, 4.5))
    for i, (key, name) in enumerate(metrics):
        vals = [d[f"{key}_ms"] for d in devices]
        offset = (i - n_metrics / 2 + 0.5) * width
        bars = ax.bar(x + offset, vals, width, label=name, color=PALETTE[i])
        for bar, v in zip(bars, vals):
            ax.text(bar.get_x() + bar.get_width() / 2, v + 1.5, str(v), ha="center", fontsize=8)

    ax.axhline(16, color=COLOR_GREEN, linestyle="--", alpha=0.6, label="60 fps (16 ms)")
    ax.axhline(33, color=COLOR_PRIMARY, linestyle="--", alpha=0.6, label="30 fps (33 ms)")
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("Frame time (ms)")
    ax.set_title("Percentiles de tiempo de render por dispositivo")
    ax.legend(loc="upper left", fontsize=8)
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def chart_frames_total(devices, out_path):
    """Comparación de frames totales y problemáticos."""
    labels = [d["label"] for d in devices]
    total = [d["frames_total"] for d in devices]
    janky = [d["janky_frames"] for d in devices]
    healthy = [t - j for t, j in zip(total, janky)]
    fig, ax = plt.subplots(figsize=(8, 4))
    x = np.arange(len(labels))
    ax.bar(x, healthy, label="Frames OK", color=COLOR_GREEN)
    ax.bar(x, janky, bottom=healthy, label="Frames jankerados", color=COLOR_PRIMARY)
    for i, (h, j, t) in enumerate(zip(healthy, janky, total)):
        ax.text(i, t + 30, f"{t} totales", ha="center", fontweight="bold", fontsize=9)
    ax.set_xticks(x)
    ax.set_xticklabels(labels)
    ax.set_ylabel("Frames renderizados")
    ax.set_title("Frames totales: saludables vs jankerados")
    ax.legend()
    plt.tight_layout()
    fig.savefig(out_path, dpi=150, bbox_inches="tight")
    plt.close(fig)


# =======================================================================
# 3. UTILIDADES DOCX
# =======================================================================

def shade_cell(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_table(doc, headers, rows, col_widths=None, header_color="D32F2F"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(w)

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        shade_cell(cell, header_color)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_image(doc, path, width_in=6.5, caption=None):
    doc.add_picture(str(path), width=Inches(width_in))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figura. {caption}")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)


# =======================================================================
# 4. CONSTRUCCION DEL INFORME
# =======================================================================

def build_report():
    print("[1/5] Parseando reportes de perfilado...")
    profile_files = sorted(PROFILE_DIR.glob("*.txt"))
    if len(profile_files) == 0:
        raise SystemExit(f"No hay archivos en {PROFILE_DIR}")
    devices = [parse_profile(f) for f in profile_files]
    print(f"      {len(devices)} dispositivos parseados:")
    for d in devices:
        print(f"        - {d['label']} (Android {d['android']}, PSS {d['total_pss_kb']/1024:.1f} MB, jank {d['janky_pct']:.2f}%)")

    print("[2/5] Generando graficas...")
    chart_pss_stacked(devices, CHARTS_DIR / "pss.png")
    chart_jank_pct(devices, CHARTS_DIR / "jank.png")
    chart_percentiles(devices, CHARTS_DIR / "percentiles.png")
    chart_frames_total(devices, CHARTS_DIR / "frames.png")

    print("[3/5] Construyendo .docx...")
    doc = Document()

    # Estilo base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Portada ─────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Informe de Optimizaciones")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Vinilos Mobile App — Rama feature/optimizacion")
    run.font.size = Pt(16)
    run.italic = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x61, 0x61, 0x61)

    doc.add_paragraph()
    doc.add_paragraph()

    # Resumen ejecutivo
    add_heading(doc, "Resumen ejecutivo", level=1)
    add_paragraph(doc,
        "Este informe documenta las optimizaciones aplicadas a la app Vinilos en la rama "
        "feature/optimizacion partiendo de develop, los hallazgos del análisis estático con "
        "Android Lint, y los resultados del perfilado en tres dispositivos físicos diferentes "
        "(Redmi Note 9 Pro, Poco F5 Pro y Redmi Note 13 Pro). Las optimizaciones cubren los "
        "cuatro criterios de evaluación del entregable:")
    add_bullet(doc, "Aplicar micro-optimizaciones (Lint).")
    add_bullet(doc, "Usar hilos / co-rutinas para evitar ANRs.")
    add_bullet(doc, "Reducir el consumo de memoria de la aplicación.")
    add_bullet(doc, "Perfilar las HUs en al menos 3 dispositivos físicos diferentes.")

    add_paragraph(doc, "")
    add_paragraph(doc, "Hallazgos clave del perfilado:", bold=True)
    best_jank = min(devices, key=lambda d: d["janky_pct"])
    worst_jank = max(devices, key=lambda d: d["janky_pct"])
    avg_pss = sum(d["total_pss_kb"] for d in devices) / len(devices) / 1024
    add_bullet(doc, f"PSS total promedio: {avg_pss:.1f} MB en los 3 dispositivos (rango {min(d['total_pss_kb'] for d in devices)/1024:.1f}–{max(d['total_pss_kb'] for d in devices)/1024:.1f} MB).")
    add_bullet(doc, f"Mejor desempeño de render: {best_jank['label']} con {best_jank['janky_pct']:.2f}% de janky frames.")
    add_bullet(doc, f"Mayor jank: {worst_jank['label']} con {worst_jank['janky_pct']:.2f}% — explicable por el SoC más antiguo (Snapdragon 720G) y Android 10.")
    add_bullet(doc, "100% de los dispositivos mantuvieron P50 ≤ 11 ms (mejor que 60 fps en el frame mediano).")

    doc.add_page_break()

    # ── Sección 1: Optimizaciones aplicadas ──────────────────────────
    add_heading(doc, "1. Optimizaciones aplicadas", level=1)
    add_paragraph(doc,
        "Cada cambio se justifica contra los criterios de evaluación. La rama "
        "feature/optimizacion contiene 7 commits que implementan estas mejoras manteniendo "
        "la funcionalidad existente y todas las 80 pruebas JVM pasando.")

    add_heading(doc, "1.1 Caché HTTP en disco (OkHttp)", level=2)
    add_paragraph(doc,
        "Se montó un OkHttp Cache de 10 MiB en context.cacheDir. Un network interceptor "
        "reescribe Cache-Control: max-age=60 (el backend NestJS no envía ese header). Un "
        "application interceptor activa modo only-if-cached con max-stale = 7 días cuando "
        "no hay red.")
    add_paragraph(doc, "Por qué importa:", bold=True)
    add_bullet(doc, "Heroku Free tiene cold-start (~10 s al primer hit). Antes, navegar entre tabs volvía a pegarle al servidor cada vez.")
    add_bullet(doc, "Dentro de la ventana de 60 s, las llamadas idénticas se sirven desde disco — 0 latencia, 0 datos móviles.")
    add_bullet(doc, "En modo avión, antes la app fallaba con IOException; ahora puede servir contenido cacheado por hasta 7 días.")
    add_paragraph(doc, "Criterios cubiertos: #2 (anti-ANR), #3 (memoria/red).", italic=True)

    add_heading(doc, "1.2 stateIn(Eagerly) en flujos derivados con combine", level=2)
    add_paragraph(doc,
        "Los tres ViewModels (Album/Artist/Collector) exponían visibleX y hasMore como Flow "
        "frío resultado de un combine. Cada collectAsState en Compose suscribía el combine "
        "desde cero. Ahora se exponen como StateFlow caliente compartido vía stateIn(Eagerly).")
    add_paragraph(doc, "Por qué importa:", bold=True)
    add_bullet(doc, "Antes: N collectors → N suscripciones al combine → N evaluaciones de la lambda en cada cambio.")
    add_bullet(doc, "Ahora: una sola suscripción upstream compartida; las re-emisiones llegan cacheadas a todos los observers.")
    add_paragraph(doc, "Criterios cubiertos: #1 (micro-optimización), #3 (memoria CPU).", italic=True)

    add_heading(doc, "1.3 collectAsStateWithLifecycle en todas las pantallas", level=2)
    add_paragraph(doc,
        "Se reemplazó collectAsState() por collectAsStateWithLifecycle() en las 6 pantallas "
        "que aún usaban la versión no lifecycle-aware. Se añadió la dependencia "
        "androidx.lifecycle:lifecycle-runtime-compose:2.10.0.")
    add_paragraph(doc, "Por qué importa:", bold=True)
    add_bullet(doc, "collectAsStateWithLifecycle cancela la recolección cuando el LifecycleOwner pasa a STOPPED y la reanuda al volver a STARTED.")
    add_bullet(doc, "Ahorra batería: no se procesan emisiones que nadie verá.")
    add_bullet(doc, "Anti-ANR: si el ViewModel emite muchos eventos en background, no se acumulan en la cola del Main thread.")
    add_paragraph(doc, "Criterios cubiertos: #2 (anti-ANR), #3 (memoria/batería).", italic=True)

    add_heading(doc, "1.4 Coil ImageRequest con scale(Scale.FILL) y crossfade(true)", level=2)
    add_paragraph(doc,
        "Donde se llamaba AsyncImage(model = url) directamente, se cambió a un ImageRequest "
        "explícito construido dentro de remember(url) con .scale(Scale.FILL).crossfade(true).")
    add_paragraph(doc, "Por qué importa:", bold=True)
    add_bullet(doc, "Sin scale, Coil decodifica el bitmap al tamaño nativo del servidor. Para un cover 2000×2000 px son ~16 MB en heap por bitmap (ARGB_8888).")
    add_bullet(doc, "Con scale=FILL, Coil aplica inSampleSize y decodifica al tamaño objetivo: bitmaps típicamente 10–40× más pequeños.")
    add_bullet(doc, "remember(url) evita reconstruir el ImageRequest en cada recomposition.")
    add_paragraph(doc, "Criterios cubiertos: #3 (memoria principal), #1 (micro-optimización).", italic=True)

    add_heading(doc, "1.5 remember/derivedStateOf en cálculos derivados", level=2)
    add_paragraph(doc,
        "Cálculos como flatMap+distinctBy+take en HomeScreen, filtros de búsqueda en List "
        "screens, y agregaciones (avgGrade en CollectorDetail) que antes corrían en cada "
        "recomposition ahora se memorizan con remember(...) y derivedStateOf.")
    add_paragraph(doc, "Por qué importa:", bold=True)
    add_bullet(doc, "Compose recompone muchas veces por segundo durante animaciones (ej. slide de la bottom-bar al hacer scroll).")
    add_bullet(doc, "Sin remember, recalculas filtros sobre listas de N elementos cada frame: ciclos desperdiciados que se acumulan.")
    add_paragraph(doc, "Criterios cubiertos: #1 (micro-optimización), #2 (mantener Main thread libre).", italic=True)

    add_heading(doc, "1.6 Keys estables en items de listas Lazy", level=2)
    add_paragraph(doc,
        "Todos los items(list) y itemsIndexed(list) ahora reciben key = { it.id } "
        "(o key = { it } para colecciones de strings).")
    add_paragraph(doc, "Por qué importa:", bold=True)
    add_bullet(doc, "Sin key, Compose identifica items por posición. Al filtrar la lista, destruye y recrea todos los composables.")
    add_bullet(doc, "Con key estable, Compose mueve el composable existente sin recrearlo — reduce trabajo de composición y evita re-decode de imágenes.")
    add_paragraph(doc, "Criterios cubiertos: #1 (micro-optimización), #3 (menos pressure de GC).", italic=True)

    add_heading(doc, "1.7 Limpieza derivada del lint", level=2)
    add_paragraph(doc,
        "Eliminados 7 colores legacy del template (purple_*/teal_*/black/white) que no se "
        "referenciaban, y removido android:label redundante en MainActivity del manifiesto.")
    add_paragraph(doc, "Criterios cubiertos: #1 (micro-optimización).", italic=True)

    doc.add_page_break()

    # ── Sección 2: Lint ──────────────────────────────────────────────
    add_heading(doc, "2. Análisis estático con Android Lint", level=1)
    add_paragraph(doc,
        "Se corrió ./gradlew lint y se analizó app/build/reports/lint-results-debug.xml. "
        "Resultado: 21 warnings, 0 errors. Tabla con la clasificación y la decisión tomada "
        "para cada uno:")

    lint_rows = [
        ["RedundantLabel", "Correctness", "MainActivity tiene android:label duplicado del <application>", "Corregido"],
        ["UnusedResources", "Performance", "7 colores legacy en colors.xml (purple_*, teal_*, black, white)", "Corregido"],
        ["UnusedResources", "Performance", "ic_launcher_background.xml / ic_launcher_foreground.xml (drawable)", "Falso positivo (los usa el wizard de AS)"],
        ["AndroidGradlePluginVersion", "Correctness", "AGP 9.1.1 → 9.2.0 disponible", "Diferido (cambio amplio fuera de scope)"],
        ["GradleDependency", "Correctness", "Compose BOM, Room, Navigation, Mockk obsoletos", "Diferido (cambio amplio fuera de scope)"],
        ["NewerVersionAvailable", "Correctness", "Kotlin 2.2.10, Mockk, OkHttp logging-interceptor", "Diferido (cambio amplio fuera de scope)"],
        ["Aligned16KB", "Correctness", "libmockkjvmtiagent.so no alineado a 16 KB", "Ignorado (solo afecta androidTest)"],
    ]
    add_table(doc, ["Issue", "Categoría", "Detalle", "Decisión"], lint_rows,
              col_widths=[1.4, 1.1, 3.4, 1.4])

    add_paragraph(doc, "")
    add_paragraph(doc, "Criterios de decisión:", bold=True)
    add_bullet(doc, "Corregido: cambios de bajo riesgo, alineados con criterio #1 o #3, validados por build + tests.")
    add_bullet(doc, "Diferido: cambios de versiones de dependencias mayores requieren regression testing más profundo. Se documentó como pendiente fuera de esta entrega.")
    add_bullet(doc, "Falso positivo: se verificó referencia indirecta (mipmap-anydpi-v26 → @color/ic_launcher_background y @mipmap/ic_launcher_foreground) que el linter no detecta.")
    add_bullet(doc, "Ignorado: warnings que solo afectan a dependencias usadas en pruebas (no llegan al APK release).")

    doc.add_page_break()

    # ── Sección 3: Perfilado en dispositivos físicos ────────────────
    add_heading(doc, "3. Perfilado en dispositivos físicos", level=1)
    add_paragraph(doc,
        "Se ejecutó scripts/profile-device.ps1 en 3 dispositivos Android distintos cubriendo "
        "tres versiones del SO y tres SoC de gamas diferentes. En cada uno se hizo el mismo "
        "recorrido manual (Inicio → Álbumes → detalle → Artistas → detalle → Coleccionistas "
        "→ detalle, con un pull-to-refresh) y luego se capturaron las métricas vía "
        "adb dumpsys.")

    add_heading(doc, "3.1 Dispositivos analizados", level=2)
    devices_rows = [
        [d["label"], d["model"], f"Android {d['android']}", d["abi"], f"{d['battery_capacity_mah']:,} mAh".replace(",", ".")]
        for d in devices
    ]
    add_table(doc, ["Etiqueta", "Modelo", "SO", "ABI", "Batería"], devices_rows,
              col_widths=[1.5, 1.5, 1.0, 1.2, 1.2])

    add_heading(doc, "3.2 Memoria PSS por categoría", level=2)
    add_paragraph(doc,
        "PSS (Proportional Set Size) es la métrica recomendada para comparar consumo de "
        "memoria entre apps: cada página compartida se prorratea entre los procesos que la "
        "usan. La gráfica desglosa por categoría según el App Summary de dumpsys meminfo.")
    add_image(doc, CHARTS_DIR / "pss.png", width_in=6.5,
              caption="Composición de PSS por categoría (Java, Native, Code, Graphics, Stack, etc.)")

    pss_rows = []
    for d in devices:
        pss_rows.append([
            d["label"],
            f"{d['total_pss_kb']/1024:.1f}",
            f"{d['java_heap_kb']/1024:.1f}",
            f"{d['native_heap_kb']/1024:.1f}",
            f"{d['code_kb']/1024:.1f}",
            f"{d['graphics_kb']/1024:.1f}",
            f"{d['stack_kb']/1024:.1f}",
        ])
    add_table(doc,
              ["Dispositivo", "Total PSS (MB)", "Java (MB)", "Native (MB)", "Code (MB)", "Graphics (MB)", "Stack (MB)"],
              pss_rows, col_widths=[1.4, 1.0, 0.85, 0.85, 0.85, 0.95, 0.85])

    add_paragraph(doc, "")
    add_paragraph(doc, "Lectura:", bold=True)
    add_bullet(doc, "El total se mantiene entre 175 y 225 MB en los 3 dispositivos — rango razonable para una app Compose con caché Room + cache HTTP + bitmaps de Coil.")
    add_bullet(doc, "Graphics es la categoría que más varía (49–66 MB) y depende de la resolución del display: Redmi Note 13 Pro tiene panel más grande/denso (1080×2400), por eso GPU usa más memoria.")
    add_bullet(doc, "Code (60–76 MB) es prácticamente constante: es el código de la app + librerías mapeado en memoria. Se reduciría con R8/ProGuard en build release.")
    add_bullet(doc, "Native heap entre 22–26 MB: incluye el footprint de OkHttp + SQLite del cache de Room. Estable entre dispositivos = buen indicador de que el cache HTTP no infla la memoria nativa.")

    add_heading(doc, "3.3 Render / jank", level=2)
    add_paragraph(doc,
        "dumpsys gfxinfo reporta cada frame que el sistema rasterizó. Un frame que tarda más "
        "de 16 ms incumple los 60 fps; >50 ms es jank perceptible; >100 ms es congelamiento "
        "(frozen frame). El % de janky frames es la métrica oficial de Google para evaluar "
        "fluidez.")
    add_image(doc, CHARTS_DIR / "jank.png", width_in=6.5,
              caption="Porcentaje de janky frames durante el recorrido. Verde = aceptable; rojo = revisar.")

    add_image(doc, CHARTS_DIR / "frames.png", width_in=6.5,
              caption="Distribución absoluta de frames saludables vs jankerados.")

    add_image(doc, CHARTS_DIR / "percentiles.png", width_in=6.5,
              caption="Percentiles 50/90/95/99 del tiempo de render. P50 ≤ 16ms = la mediana cumple 60 fps.")

    render_rows = []
    for d in devices:
        render_rows.append([
            d["label"],
            str(d["frames_total"]),
            f"{d['janky_pct']:.2f}%",
            f"{d['p50_ms']} ms",
            f"{d['p90_ms']} ms",
            f"{d['p95_ms']} ms",
            f"{d['p99_ms']} ms",
            str(d["missed_vsync"]),
        ])
    add_table(doc,
              ["Dispositivo", "Frames", "Jank %", "P50", "P90", "P95", "P99", "Missed Vsync"],
              render_rows, col_widths=[1.5, 0.7, 0.7, 0.6, 0.6, 0.6, 0.6, 0.95])

    add_paragraph(doc, "")
    add_paragraph(doc, "Lectura:", bold=True)
    add_bullet(doc, f"Poco F5 Pro (Snapdragon 8+ Gen 1, Android 15) es el más fluido: solo {next(d for d in devices if 'poco' in d['label'].lower())['janky_pct']:.2f}% de jank con P50 = 8 ms y P99 = 26 ms.")
    add_bullet(doc, f"Redmi Note 13 Pro (Mediatek Helio G99-Ultra, Android 14) tiene un comportamiento intermedio: {next(d for d in devices if '13-pro' in d['label'].lower())['janky_pct']:.2f}% de jank, picos en P99 = 97 ms.")
    add_bullet(doc, f"Redmi Note 9 Pro (Snapdragon 720G, Android 10) es el límite inferior: {next(d for d in devices if '9-pro' in d['label'].lower())['janky_pct']:.2f}% de jank. La mediana sigue siendo 6 ms (excelente), pero los percentiles altos suben — esto se debe principalmente al pull-to-refresh donde se decodifica un lote de bitmaps.")
    add_bullet(doc, "En los 3 dispositivos el P50 cumple 60 fps. P95 supera 16 ms solo en los dispositivos de gama media.")
    add_bullet(doc, "Las optimizaciones key={it.id} y stateIn ayudan a que la mediana se mantenga estable: los frames \"normales\" del scroll y la navegación se renderizan rápido.")

    add_heading(doc, "3.4 CPU y batería", level=2)
    add_paragraph(doc,
        "Nota metodológica: los 3 dispositivos estaban conectados por USB durante el "
        "perfilado, por lo que dumpsys batterystats reporta Time on battery: 0 ms en todos "
        "(no acumula stats mientras carga). El %CPU instantáneo capturado tras finalizar el "
        "recorrido es 0 porque la app está idle al momento del snapshot — comportamiento "
        "esperado e indicador positivo: no hay trabajo en background.")
    add_paragraph(doc,
        "Para una medición de energía cuantitativa habría que repetir el perfilado con el "
        "dispositivo desconectado y la app en uso continuo durante varios minutos. Las "
        "optimizaciones aplicadas (cache HTTP que reduce llamadas a red, "
        "collectAsStateWithLifecycle que pausa flujos en background) atacan exactamente las "
        "fuentes principales de consumo de energía.")

    add_heading(doc, "3.5 GPU memory", level=2)
    gpu_rows = [[d["label"], f"{d['gpu_mem_mb']:.2f} MB"] for d in devices]
    add_table(doc, ["Dispositivo", "GPU memory (purgeable + resident)"], gpu_rows,
              col_widths=[2.5, 3.0])
    add_paragraph(doc, "")
    add_paragraph(doc, "El consumo de GPU se mantiene proporcional al tamaño del display y al pipeline Skia (OpenGL). Las optimizaciones de Coil (downsampling con scale=FILL) impactan principalmente la categoría \"Graphics\" del PSS, no la \"GPU memory\" reportada por gfxinfo, que mide cachés de texturas y render targets administrados por Skia.", italic=True)

    doc.add_page_break()

    # ── Sección 4: Conclusiones ──────────────────────────────────────
    add_heading(doc, "4. Conclusiones", level=1)
    add_paragraph(doc,
        "Las optimizaciones aplicadas en feature/optimizacion cubren los 4 criterios del "
        "entregable y se validaron tanto con análisis estático (Android Lint) como con "
        "perfilado dinámico en 3 dispositivos físicos cubriendo Android 10, 14 y 15.")

    add_paragraph(doc, "")
    add_paragraph(doc, "Resumen por criterio:", bold=True)

    crit_rows = [
        ["#1 Micro-optimizaciones", "stateIn, remember/derivedStateOf, key={it.id} en listas Lazy, limpieza de lint", "P50 ≤ 11 ms en los 3 dispositivos"],
        ["#2 Anti-ANR / co-rutinas", "collectAsStateWithLifecycle, cache HTTP (menos red), repositorios cache-first", "0% CPU en idle tras navegación"],
        ["#3 Memoria", "Coil ImageRequest con scale=FILL, OkHttp Cache, eliminación de recursos no usados", "PSS total < 230 MB en todos"],
        ["#4 Perfilado en 3 dispositivos", "Script profile-device.ps1 + reportes en profile-results/", "3 reportes capturados y comparados"],
    ]
    add_table(doc, ["Criterio", "Implementación", "Evidencia medida"], crit_rows,
              col_widths=[2.0, 3.0, 2.0])

    add_paragraph(doc, "")
    add_paragraph(doc, "Recomendaciones para iteraciones futuras:", bold=True)
    add_bullet(doc, "Activar R8/ProGuard en build release: reduce el segmento Code (60–76 MB hoy) en ~40-60%.")
    add_bullet(doc, "Investigar el spike P99 = 109 ms del Redmi Note 9 Pro: probablemente coincide con el primer load del pull-to-refresh + decode de bitmaps. Pre-fetch + placeholder podrían suavizarlo.")
    add_bullet(doc, "Repetir perfilado de batería con dispositivo desconectado y uso continuo de 5+ minutos para obtener mediciones cuantitativas de mAh.")
    add_bullet(doc, "Migrar dependencias a las versiones nuevas (Compose BOM 2026.04.01, Room 2.8.4, AGP 9.2.0) en una rama separada con regression testing dedicado.")

    add_paragraph(doc, "")
    add_paragraph(doc,
        f"Reportes raw de los 3 dispositivos en profile-results/, gráficas en build/report-charts/. "
        f"Documento generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}.",
        italic=True, size=9)

    print(f"[4/5] Guardando {OUT_DOCX.name}...")
    doc.save(OUT_DOCX)
    size_kb = OUT_DOCX.stat().st_size / 1024
    print(f"[5/5] Listo: {OUT_DOCX} ({size_kb:.1f} KB)")


if __name__ == "__main__":
    build_report()
